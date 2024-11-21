import os
from docx import Document
from docx.shared import Inches
from colorama import Fore, Style, init
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Initialize colorma
init()

def create_image_docx(image_folder, output_file):
    # Create a new document
    doc = Document()

    # Set the margins to the minimum (no margin space)
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.2)
        section.right_margin = Inches(0.2)
        section.top_margin = Inches(0.2)
        section.bottom_margin = Inches(0.2)

    # Get the list of image files from the specified folder
    image_files = [f for f in os.listdir(image_folder) if f.endswith(('jpeg', 'png', 'jpg', 'gif', 'bmp'))]
    
    # Sort the images alphabetically or in any order you prefer
    image_files.sort()

    # Initialize counter for images per row
    row = doc.add_paragraph()  # A new paragraph for the images on the current line
    count = 0

    # Loop through images and add them to the document
    for index, image_file in enumerate(image_files):
        image_path = os.path.join(image_folder, image_file)

        # Add image to the current paragraph (row)
        run = row.add_run(' ')
        run.add_picture(image_path, width=Inches(2.5), height=Inches(3.5))  # Resize to the desired dimensions





    # Save the document
    doc.save(output_file)
    print(f"Document saved as {output_file}")

if __name__ == "__main__":
    # Get the folder containing images from the user
    image_folder = 'mtg_images'
    
    # Prompt for the output file name
    output_file = input("Enter the name for the output .docx file: ") + '.docx'
    
    # Call the function to create the document
    create_image_docx(image_folder, output_file)

